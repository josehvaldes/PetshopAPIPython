# This script converts .docx and .pdf files in a specified folder into a JSONL format suitable for training a language model.
import re
import os
import json
import fitz  # PyMuPDF
from docx import Document
from tqdm import tqdm
from docx import Document


def extract_clean_text(docx_path):
    doc = Document(docx_path)
    full_text = []
    print(f"Extracting text from {docx_path}...")
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_text.append(cell_text)

    return "\n".join(full_text)

def extract_doc_text(path):
    doc = Document(path)
    print(f"Extracting text from {path}...")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"First paragraph: {doc.paragraphs[0].text if doc.paragraphs else 'N/A'}")
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_pdf_text(path):
    doc = fitz.open(path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def normalize_text(text):
    text = re.sub(r'\s+', ' ', text)  # Collapse multiple spaces/newlines
    text = re.sub(r'[^\S\r\n]{2,}', ' ', text)  # Remove excessive indentation
    return text.strip()

def create_jsonl_from_folder(folder_path, output_path):
    entries = []
    for filename in tqdm(os.listdir(folder_path)):
        
        print(f"Processing {filename}...")
        full_path = os.path.join(folder_path, filename)
        if filename.endswith(".docx") :
            text = extract_clean_text(full_path)
            text = normalize_text(text)
        elif filename.endswith(".pdf"):
            text = extract_pdf_text(full_path)
        else:
            continue

        if len(text.strip()) < 100:
            print(f"Skipped (too short) : {len(text.strip())}")
            continue  # Skip short or empty docs

        entry = {
            "instruction": "Resume el siguiente documento sobre perros.",
            "input": text.strip(),
            "output": ""  # unsupervised training
        }
        entries.append(entry)

    with open(output_path, "w", encoding="utf-8") as f:
        for item in entries:
            f.write(json.dumps(item, ensure_ascii=False) + "\n")

    print(f"✅ Saved {len(entries)} entries to {output_path}")

dogsfiles_path = "c:/personal/_gemma/customdocs/raw/dogs_breeds_docx/"
dogs_dataset_jsonl = "c:/personal/_gemma/customdocs/converted/dogs_dataset.jsonl"
# Example usage
create_jsonl_from_folder(dogsfiles_path, dogs_dataset_jsonl)